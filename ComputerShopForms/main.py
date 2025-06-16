# Import required modules
import datetime
from diagnosticform import collect_diagnostic_info
from invoiceform import collect_invoice_info
from loggingform import collect_logging_info
from docx import Document

def save_repair_ticket():
    try:
        # Create a new Word document
        document = Document()
        document.add_heading('Repair Ticket', level=1)

        # Collect diagnostic info
        diagnostic_data = collect_diagnostic_info()
        document.add_heading('Diagnostic Form', level=2)
        for key, value in diagnostic_data.items():
            document.add_paragraph(f"{key}: {value}")

        # Collect invoice info
        invoice_data = collect_invoice_info(diagnostic_data["Issue Type"], diagnostic_data["Description"])
        document.add_heading('Invoice Form', level=2)
        for key, value in invoice_data.items():
            document.add_paragraph(f"{key}: {value}")

        # Ask if logging form should be started
        start_logging = input("Would you like to begin the logging form? (yes/no): ").strip().lower()
        if start_logging == 'yes':
            logging_data = collect_logging_info(invoice_data["Urgency"])
            document.add_heading('Logging Form', level=2)
            for key, value in logging_data.items():
                document.add_paragraph(f"{key}: {value}")

        # Save document with timestamp
        filename = f"Repair_Ticket_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(filename)
        print(f"Repair ticket saved as {filename}")
    # Print an error message if an exception occurs
    except Exception as e:
        print(f"An error occurred while creating the repair ticket: {e}")
# 
if __name__ == "__main__":
    save_repair_ticket()